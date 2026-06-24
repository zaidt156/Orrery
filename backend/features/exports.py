from __future__ import annotations

import asyncio
import csv
import datetime
import html
import io
import json
import re
import uuid
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from markdown_it import MarkdownIt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    HRFlowable,
    LongTable,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    TableStyle,
)
from sqlalchemy import select

from backend.core.database import get_sessionmaker
from backend.core.models import Conversation, Message

MAX_EXPORT_CHARS = 2_000_000
MAX_TABLE_CELLS = 50_000
_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
_SLUG_CHARS = re.compile(r"[^a-z0-9]+")
_FORMULA_PREFIXES = ("=", "+", "-", "@")
_MARKDOWN = MarkdownIt("commonmark", {"html": False, "linkify": False}).enable("table")
SUPPORTED_FORMATS = {"pdf", "docx", "xlsx", "pptx", "csv", "md", "txt", "html", "json"}
TEXT_FORMATS = {"csv", "md", "txt", "html", "json"}

_FENCE_RE = re.compile(r"```([a-zA-Z0-9_+.#-]*)[^\n]*\n([\s\S]*?)```")
_TABLE_LINE_RE = re.compile(r"^\s*\|.*\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
_CHATTER_PREFIX_RE = re.compile(
    r"^\s*(sure|okay|ok|here(?:'s| is)|i (?:created|made|prepared|generated)|below is|"
    r"you can (?:download|save)|download (?:it|the)|note:)\b",
    re.IGNORECASE,
)
_CHATTER_TRAILING_RE = re.compile(
    r"^\s*(let me know|you can ask|i can also|if you want|hope this helps)\b",
    re.IGNORECASE,
)

_FORMAT_PATTERNS: tuple[tuple[str, tuple[re.Pattern[str], ...]], ...] = (
    ("pdf", (re.compile(r"\bpdf\b|\.pdf\b", re.I), re.compile(r"\breport\b", re.I))),
    ("docx", (re.compile(r"\b(word|docx?|document)\b|\.docx?\b", re.I),)),
    ("xlsx", (re.compile(r"\b(excel|xlsx?|spreadsheet|workbook|sheet)\b|\.xlsx?\b", re.I),)),
    ("pptx", (re.compile(r"\b(powerpoint|pptx?|presentation|slide\s*deck|slides?)\b|\.pptx?\b", re.I),)),
    ("csv", (re.compile(r"\bcsv\b|\.csv\b", re.I),)),
    ("json", (re.compile(r"\bjson\b|\.json\b", re.I),)),
    ("html", (re.compile(r"\bhtml\b|web\s?page|\.html?\b", re.I),)),
    ("md", (re.compile(r"\bmarkdown\b|\.md\b", re.I),)),
    ("txt", (re.compile(r"\b(?:plain\s+text|text file|txt)\b|\.txt\b", re.I),)),
)


class ExportNotFound(Exception):
    pass


class ExportTooLarge(Exception):
    pass


@dataclass(frozen=True)
class InlineSpan:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


@dataclass
class ExportBlock:
    kind: str
    spans: list[InlineSpan] = field(default_factory=list)
    text: str = ""
    level: int = 0
    language: str = ""
    rows: list[list[str]] = field(default_factory=list)


@dataclass(frozen=True)
class ExportResult:
    content: bytes
    media_type: str
    filename: str


def _clean_text(value: str) -> str:
    return _CONTROL_CHARS.sub("", str(value or "")).replace("\r\n", "\n").replace("\r", "\n")


def _plain_spans(spans: list[InlineSpan]) -> str:
    return "".join(span.text for span in spans)


def requested_formats(prompt: str) -> list[str]:
    """Return safe export formats explicitly requested by the user prompt."""
    text = _clean_text(prompt).lower()
    if not text:
        return []
    found: list[str] = []
    for export_format, patterns in _FORMAT_PATTERNS:
        if any(pattern.search(text) for pattern in patterns):
            found.append(export_format)
    return found


def _fenced_blocks(content: str, languages: set[str] | None = None) -> list[str]:
    blocks: list[str] = []
    for match in _FENCE_RE.finditer(content or ""):
        lang = (match.group(1) or "").lower().strip()
        lang = {"js": "javascript", "ts": "typescript", "py": "python"}.get(lang, lang)
        if languages is None or lang in languages:
            blocks.append(_clean_text(match.group(2)).strip())
    return blocks


def _looks_structural(line: str) -> bool:
    stripped = line.strip()
    return bool(
        stripped.startswith(("#", "|", "- ", "* ", "1. ", "{", "[", "<", "```"))
        or re.match(r"^\d+[.)]\s+", stripped)
    )


def strip_chat_wrappers(content: str) -> str:
    """Remove common assistant wrapper lines while preserving the requested artifact body."""
    lines = _clean_text(content).strip().split("\n")
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()

    # Drop leading helper chatter only when the real artifact starts after it.
    while lines and _CHATTER_PREFIX_RE.match(lines[0]) and any(_looks_structural(line) for line in lines[1:6]):
        lines.pop(0)
        while lines and not lines[0].strip():
            lines.pop(0)

    # Drop common closing offers that should not enter generated files.
    while lines and _CHATTER_TRAILING_RE.match(lines[-1]):
        lines.pop()
        while lines and not lines[-1].strip():
            lines.pop()
    return "\n".join(lines).strip()


def _extract_tables(markdown: str) -> str:
    lines = _clean_text(markdown).split("\n")
    tables: list[list[str]] = []
    i = 0
    while i < len(lines):
        if i + 1 < len(lines) and _TABLE_LINE_RE.match(lines[i]) and _TABLE_SEP_RE.match(lines[i + 1]):
            block = [lines[i], lines[i + 1]]
            i += 2
            while i < len(lines) and _TABLE_LINE_RE.match(lines[i]):
                block.append(lines[i])
                i += 1
            tables.append(block)
            continue
        i += 1
    return "\n\n".join("\n".join(table) for table in tables).strip()


def _extract_json(content: str) -> str:
    for block in _fenced_blocks(content, {"json"}):
        try:
            return json.dumps(json.loads(block), indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            pass
    cleaned = strip_chat_wrappers(content)
    try:
        return json.dumps(json.loads(cleaned), indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return json.dumps({"content": blocks_to_plain(parse_markdown(cleaned))}, indent=2, ensure_ascii=False)


def select_export_content(prompt: str, content: str, export_format: str) -> str:
    """Extract the exact file body from a saved reply for the requested format."""
    cleaned = strip_chat_wrappers(content)
    if export_format == "json":
        return _extract_json(content)
    if export_format == "html":
        blocks = _fenced_blocks(content, {"html"})
        if blocks:
            return blocks[0]
        if re.search(r"<!doctype html|<html[\s>]", cleaned, re.I):
            return cleaned
        return cleaned
    if export_format in {"md", "txt"}:
        markdown_blocks = _fenced_blocks(content, {"markdown", "md", "text", "txt", "plain"})
        return markdown_blocks[0] if markdown_blocks else cleaned
    if export_format in {"xlsx", "csv"}:
        tables = _extract_tables(cleaned)
        return tables or cleaned
    return cleaned


def blocks_to_plain(blocks: list[ExportBlock]) -> str:
    parts: list[str] = []
    for block in blocks:
        if block.kind == "heading":
            parts.append(_plain_spans(block.spans))
        elif block.kind in {"paragraph", "quote", "bullet", "number"}:
            prefix = "- " if block.kind == "bullet" else (f"{block.text}. " if block.kind == "number" else "")
            parts.append(prefix + _plain_spans(block.spans))
        elif block.kind == "code":
            parts.append(block.text)
        elif block.kind == "table":
            for row in block.rows:
                parts.append("\t".join(row))
        elif block.kind == "rule":
            parts.append("---")
    return "\n".join(part for part in parts if part.strip()).strip()


def _inline_spans(token) -> list[InlineSpan]:
    children = token.children or []
    if not children:
        return [InlineSpan(_clean_text(token.content))] if token.content else []

    spans: list[InlineSpan] = []
    bold = False
    italic = False

    def add(text: str, *, code: bool = False) -> None:
        text = _clean_text(text)
        if not text:
            return
        span = InlineSpan(text, bold=bold, italic=italic, code=code)
        if spans and spans[-1].bold == span.bold and spans[-1].italic == span.italic and spans[-1].code == span.code:
            previous = spans[-1]
            spans[-1] = InlineSpan(previous.text + span.text, span.bold, span.italic, span.code)
        else:
            spans.append(span)

    for child in children:
        if child.type == "strong_open":
            bold = True
        elif child.type == "strong_close":
            bold = False
        elif child.type == "em_open":
            italic = True
        elif child.type == "em_close":
            italic = False
        elif child.type == "text":
            add(child.content)
        elif child.type == "code_inline":
            add(child.content, code=True)
        elif child.type in {"softbreak", "hardbreak"}:
            add("\n")
        elif child.type == "image":
            add(f"[image: {child.content or 'image'}]")
    return spans


def _parse_table(tokens: list, start: int) -> tuple[ExportBlock, int]:
    rows: list[list[str]] = []
    row: list[str] | None = None
    cell = ""
    i = start + 1
    while i < len(tokens):
        token = tokens[i]
        if token.type == "table_close":
            break
        if token.type == "tr_open":
            row = []
        elif token.type in {"th_open", "td_open"}:
            cell = ""
        elif token.type == "inline" and row is not None:
            cell = _plain_spans(_inline_spans(token))
        elif token.type in {"th_close", "td_close"} and row is not None:
            row.append(cell)
        elif token.type == "tr_close" and row is not None:
            rows.append(row)
            row = None
        i += 1

    cell_count = sum(len(row) for row in rows)
    if cell_count > MAX_TABLE_CELLS:
        raise ExportTooLarge("The reply contains a table that is too large to export safely.")
    return ExportBlock(kind="table", rows=rows), i


def parse_markdown(markdown: str) -> list[ExportBlock]:
    text = _clean_text(markdown)
    if len(text) > MAX_EXPORT_CHARS:
        raise ExportTooLarge("This reply is too large to export safely.")

    tokens = _MARKDOWN.parse(text)
    blocks: list[ExportBlock] = []
    lists: list[dict] = []
    current_item: tuple[str, int, int] | None = None
    quote_depth = 0
    i = 0

    while i < len(tokens):
        token = tokens[i]
        if token.type == "table_open":
            table, i = _parse_table(tokens, i)
            blocks.append(table)
        elif token.type == "bullet_list_open":
            lists.append({"kind": "bullet", "counter": 0})
        elif token.type == "ordered_list_open":
            start = int(token.attrGet("start") or 1)
            lists.append({"kind": "number", "counter": start - 1})
        elif token.type in {"bullet_list_close", "ordered_list_close"}:
            if lists:
                lists.pop()
        elif token.type == "list_item_open" and lists:
            active = lists[-1]
            if active["kind"] == "number":
                active["counter"] += 1
            current_item = (active["kind"], active["counter"], max(0, len(lists) - 1))
        elif token.type == "list_item_close":
            current_item = None
        elif token.type == "blockquote_open":
            quote_depth += 1
        elif token.type == "blockquote_close":
            quote_depth = max(0, quote_depth - 1)
        elif token.type == "heading_open" and i + 1 < len(tokens):
            inline = tokens[i + 1]
            blocks.append(
                ExportBlock(
                    kind="heading",
                    spans=_inline_spans(inline),
                    level=int(token.tag[1:]) if token.tag[1:].isdigit() else 2,
                )
            )
        elif token.type == "paragraph_open" and i + 1 < len(tokens):
            inline = tokens[i + 1]
            if inline.type == "inline":
                if current_item:
                    kind, number, depth = current_item
                    blocks.append(
                        ExportBlock(
                            kind=kind,
                            spans=_inline_spans(inline),
                            level=depth,
                            text=str(number) if kind == "number" else "",
                        )
                    )
                else:
                    blocks.append(
                        ExportBlock(
                            kind="quote" if quote_depth else "paragraph",
                            spans=_inline_spans(inline),
                            level=max(0, quote_depth - 1),
                        )
                    )
        elif token.type in {"fence", "code_block"}:
            blocks.append(
                ExportBlock(
                    kind="code",
                    text=_clean_text(token.content).rstrip("\n"),
                    language=_clean_text((token.info or "").split()[0]),
                )
            )
        elif token.type == "hr":
            blocks.append(ExportBlock(kind="rule"))
        i += 1

    return blocks


def _slug(value: str) -> str:
    slug = _SLUG_CHARS.sub("-", value.lower()).strip("-")
    return (slug[:60] or "orrery-reply").rstrip("-")


def _pdf_fonts() -> tuple[str, str]:
    sans = "Helvetica"
    mono = "Courier"
    candidates = [
        (
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
            Path("C:/Windows/Fonts/ariali.ttf"),
            Path("C:/Windows/Fonts/arialbi.ttf"),
        ),
    ]
    for regular, bold, italic, bold_italic in candidates:
        if all(path.is_file() for path in (regular, bold, italic, bold_italic)):
            try:
                pdfmetrics.registerFont(TTFont("OrrerySans", str(regular)))
                pdfmetrics.registerFont(TTFont("OrrerySans-Bold", str(bold)))
                pdfmetrics.registerFont(TTFont("OrrerySans-Italic", str(italic)))
                pdfmetrics.registerFont(TTFont("OrrerySans-BoldItalic", str(bold_italic)))
                pdfmetrics.registerFontFamily(
                    "OrrerySans",
                    normal="OrrerySans",
                    bold="OrrerySans-Bold",
                    italic="OrrerySans-Italic",
                    boldItalic="OrrerySans-BoldItalic",
                )
                sans = "OrrerySans"
                break
            except Exception:
                pass
    return sans, mono


def _pdf_markup(spans: list[InlineSpan], sans: str, mono: str) -> str:
    parts: list[str] = []
    for span in spans:
        value = html.escape(span.text).replace("\n", "<br/>")
        if span.code:
            value = f'<font name="{mono}" color="#244b7a">{value}</font>'
        if span.italic:
            value = f"<i>{value}</i>"
        if span.bold:
            value = f"<b>{value}</b>"
        parts.append(value)
    return "".join(parts) or " "


def build_pdf(title: str, model: str, blocks: list[ExportBlock]) -> bytes:
    output = io.BytesIO()
    sans, mono = _pdf_fonts()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "OrreryBody",
        parent=styles["BodyText"],
        fontName=sans,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#172033"),
        spaceAfter=7,
    )
    heading_styles = {
        level: ParagraphStyle(
            f"OrreryHeading{level}",
            parent=body,
            fontName=sans,
            fontSize=max(12, 21 - level * 2),
            leading=max(15, 24 - level * 2),
            textColor=colors.HexColor("#0b1020"),
            spaceBefore=10,
            spaceAfter=6,
        )
        for level in range(1, 7)
    }
    code_style = ParagraphStyle(
        "OrreryCode",
        parent=body,
        fontName=mono,
        fontSize=8.4,
        leading=11,
        backColor=colors.HexColor("#f3f5f9"),
        borderColor=colors.HexColor("#d7deea"),
        borderWidth=0.5,
        borderPadding=7,
        spaceAfter=9,
    )
    quote_style = ParagraphStyle(
        "OrreryQuote",
        parent=body,
        leftIndent=12,
        borderColor=colors.HexColor("#8fa5c8"),
        borderWidth=1.5,
        borderPadding=7,
        textColor=colors.HexColor("#4f5d73"),
    )
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title,
        author="Orrery",
    )
    story = [
        Paragraph(html.escape(title), heading_styles[1]),
        Paragraph(f"Model: {html.escape(model or 'unknown')}", ParagraphStyle(
            "OrreryMeta", parent=body, fontSize=8.5, textColor=colors.HexColor("#69758a")
        )),
        Spacer(1, 4 * mm),
    ]

    for block in blocks:
        if block.kind == "heading":
            story.append(Paragraph(_pdf_markup(block.spans, sans, mono), heading_styles[min(6, max(1, block.level))]))
        elif block.kind == "paragraph":
            story.append(Paragraph(_pdf_markup(block.spans, sans, mono), body))
        elif block.kind == "quote":
            story.append(Paragraph(_pdf_markup(block.spans, sans, mono), quote_style))
        elif block.kind in {"bullet", "number"}:
            prefix = "•" if block.kind == "bullet" else f"{block.text}."
            list_style = ParagraphStyle(
                f"OrreryList{block.level}",
                parent=body,
                leftIndent=12 + block.level * 12,
                firstLineIndent=-10,
            )
            story.append(Paragraph(f"{prefix} {_pdf_markup(block.spans, sans, mono)}", list_style))
        elif block.kind == "code":
            if block.language:
                story.append(Paragraph(html.escape(block.language.upper()), ParagraphStyle(
                    "OrreryCodeLabel", parent=body, fontName=mono, fontSize=7.5,
                    textColor=colors.HexColor("#69758a"), spaceAfter=2
                )))
            story.append(Preformatted(block.text or " ", code_style))
        elif block.kind == "rule":
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cdd5e2"), spaceBefore=5, spaceAfter=8))
        elif block.kind == "table" and block.rows:
            width = max(len(row) for row in block.rows)
            normalized = [row + [""] * (width - len(row)) for row in block.rows]
            data = [
                [Paragraph(html.escape(cell), body) for cell in row]
                for row in normalized
            ]
            table = LongTable(data, colWidths=[doc.width / width] * width, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8edf6")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#172033")),
                ("FONTNAME", (0, 0), (-1, 0), sans),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c6cfdd")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.extend([table, Spacer(1, 3 * mm)])

    def footer(canvas, document) -> None:
        canvas.saveState()
        canvas.setFont(sans, 8)
        canvas.setFillColor(colors.HexColor("#69758a"))
        canvas.drawRightString(A4[0] - 18 * mm, 10 * mm, f"Orrery · {document.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return output.getvalue()


def _docx_runs(paragraph, spans: list[InlineSpan]) -> None:
    for span in spans:
        run = paragraph.add_run(span.text)
        run.bold = span.bold
        run.italic = span.italic
        if span.code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def _shade_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    properties.append(shade)


def build_docx(title: str, model: str, blocks: list[ExportBlock]) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    document.core_properties.title = title
    document.core_properties.author = "Orrery"
    document.add_heading(title, level=0)
    meta = document.add_paragraph()
    meta_run = meta.add_run(f"Model: {model or 'unknown'}")
    meta_run.italic = True
    meta_run.font.size = Pt(9)

    for block in blocks:
        if block.kind == "heading":
            paragraph = document.add_heading(level=min(9, max(1, block.level)))
            _docx_runs(paragraph, block.spans)
        elif block.kind in {"paragraph", "quote"}:
            paragraph = document.add_paragraph()
            _docx_runs(paragraph, block.spans)
            if block.kind == "quote":
                paragraph.paragraph_format.left_indent = Inches(0.25)
                _shade_paragraph(paragraph, "EEF2F8")
        elif block.kind in {"bullet", "number"}:
            style = "List Bullet" if block.kind == "bullet" else "List Number"
            paragraph = document.add_paragraph(style=style)
            paragraph.paragraph_format.left_indent = Inches(0.25 + block.level * 0.2)
            _docx_runs(paragraph, block.spans)
        elif block.kind == "code":
            if block.language:
                label = document.add_paragraph()
                label_run = label.add_run(block.language.upper())
                label_run.bold = True
                label_run.font.name = "Consolas"
                label_run.font.size = Pt(8)
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block.text or " ")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            _shade_paragraph(paragraph, "F1F3F6")
        elif block.kind == "rule":
            paragraph = document.add_paragraph()
            border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:color"), "C9D2E0")
            border.append(bottom)
            paragraph._p.get_or_add_pPr().append(border)
        elif block.kind == "table" and block.rows:
            width = max(len(row) for row in block.rows)
            table = document.add_table(rows=len(block.rows), cols=width)
            table.style = "Table Grid"
            for row_index, row in enumerate(block.rows):
                for column_index in range(width):
                    value = row[column_index] if column_index < len(row) else ""
                    cell = table.cell(row_index, column_index)
                    cell.text = value
                    if row_index == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
            document.add_paragraph()

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _excel_safe(value: str) -> str:
    text = _clean_text(value)
    if text.lstrip().startswith(_FORMULA_PREFIXES):
        text = "'" + text
    return text


def _excel_chunks(value: str, size: int = 32_000) -> list[str]:
    text = _excel_safe(value)
    return [text[i:i + size] for i in range(0, len(text), size)] or [""]


def build_xlsx(title: str, model: str, blocks: list[ExportBlock]) -> bytes:
    workbook = Workbook()
    reply = workbook.active
    reply.title = "Reply"
    reply.freeze_panes = "A5"
    reply.column_dimensions["A"].width = 18
    reply.column_dimensions["B"].width = 100
    reply["A1"] = "Orrery reply"
    reply["A1"].font = Font(size=16, bold=True, color="172033")
    reply["A2"] = "Title"
    reply["B2"] = _excel_safe(title)
    reply["A3"] = "Model"
    reply["B3"] = _excel_safe(model or "unknown")
    reply["A4"] = "Type"
    reply["B4"] = "Content"
    for cell in reply[4]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="26314F")

    table_number = 0
    row_number = 5
    for block in blocks:
        if block.kind == "table":
            table_number += 1
            reply.cell(row_number, 1, "Table")
            reply.cell(row_number, 2, f"See worksheet Table {table_number}")
            row_number += 1
            if not block.rows:
                continue
            sheet = workbook.create_sheet(f"Table {table_number}")
            width = max(len(row) for row in block.rows)
            for row_index, row in enumerate(block.rows, start=1):
                for column_index in range(1, width + 1):
                    value = row[column_index - 1] if column_index <= len(row) else ""
                    cell = sheet.cell(row_index, column_index, _excel_safe(value)[:32_767])
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                    if row_index == 1:
                        cell.font = Font(bold=True, color="FFFFFF")
                        cell.fill = PatternFill("solid", fgColor="26314F")
            sheet.freeze_panes = "A2"
            if len(block.rows) > 1:
                sheet.auto_filter.ref = f"A1:{get_column_letter(width)}{len(block.rows)}"
            for column_index in range(1, width + 1):
                values = [
                    len(str(sheet.cell(row_index, column_index).value or ""))
                    for row_index in range(1, min(len(block.rows), 200) + 1)
                ]
                sheet.column_dimensions[get_column_letter(column_index)].width = min(48, max([10, *values]) + 2)
            continue

        content = block.text if block.kind == "code" else _plain_spans(block.spans)
        label = {
            "heading": f"Heading {block.level}",
            "paragraph": "Paragraph",
            "quote": "Quote",
            "bullet": "Bullet",
            "number": "Numbered item",
            "code": f"Code ({block.language or 'plain'})",
            "rule": "Divider",
        }.get(block.kind, block.kind.title())
        for part in _excel_chunks(content):
            reply.cell(row_number, 1, label)
            reply.cell(row_number, 2, part)
            reply.cell(row_number, 2).alignment = Alignment(vertical="top", wrap_text=True)
            row_number += 1

    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()



def build_csv(blocks: list[ExportBlock]) -> bytes:
    output = io.StringIO(newline="")
    writer = csv.writer(output)
    tables = [block for block in blocks if block.kind == "table" and block.rows]
    if tables:
        for index, table in enumerate(tables):
            if index:
                writer.writerow([])
            for row in table.rows:
                writer.writerow([_excel_safe(cell) for cell in row])
    else:
        writer.writerow(["Type", "Content"])
        for block in blocks:
            if block.kind == "rule":
                continue
            label = block.kind.title()
            content = block.text if block.kind == "code" else _plain_spans(block.spans)
            if content.strip():
                writer.writerow([label, _excel_safe(content)])
    return output.getvalue().encode("utf-8-sig")


def build_html(title: str, model: str, content: str) -> bytes:
    if re.search(r"<!doctype html|<html[\s>]", content, re.I):
        return _clean_text(content).encode("utf-8")
    body = _MARKDOWN.render(_clean_text(content))
    document = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    :root {{ color-scheme: light; font-family: Arial, sans-serif; color: #172033; }}
    body {{ max-width: 880px; margin: 32px auto; padding: 0 24px 48px; line-height: 1.55; }}
    h1, h2, h3 {{ color: #0b1020; line-height: 1.25; }}
    table {{ border-collapse: collapse; width: 100%; margin: 14px 0; }}
    th, td {{ border: 1px solid #c6cfdd; padding: 7px 9px; text-align: left; vertical-align: top; }}
    th {{ background: #e8edf6; }}
    pre {{ background: #f3f5f9; border: 1px solid #d7deea; border-radius: 8px; padding: 12px; overflow-x: auto; }}
    code {{ font-family: Consolas, monospace; }}
    blockquote {{ border-left: 4px solid #8fa5c8; margin-left: 0; padding-left: 14px; color: #4f5d73; }}
    .meta {{ color: #69758a; font-size: 12px; margin-bottom: 24px; }}
  </style>
</head>
<body>
  <h1>{html.escape(title)}</h1>
  <div class="meta">Model: {html.escape(model or 'unknown')}</div>
  {body}
</body>
</html>"""
    return document.encode("utf-8")


def _plain_or_markdown_text(content: str, blocks: list[ExportBlock], markdown: bool) -> bytes:
    value = _clean_text(content if markdown else blocks_to_plain(blocks))
    return (value.strip() + "\n").encode("utf-8")


def _xml(value: str) -> str:
    return html.escape(str(value or ""), quote=True)


def _slide_text_from_block(block: ExportBlock) -> list[str]:
    if block.kind == "heading":
        return []
    if block.kind == "paragraph":
        return [_plain_spans(block.spans)]
    if block.kind == "quote":
        return [f"Quote: {_plain_spans(block.spans)}"]
    if block.kind == "bullet":
        return [f"- {_plain_spans(block.spans)}"]
    if block.kind == "number":
        return [f"{block.text}. {_plain_spans(block.spans)}"]
    if block.kind == "code":
        lines = [line for line in block.text.splitlines() if line.strip()]
        label = f"Code ({block.language or 'text'})"
        return [label, *lines[:10]]
    if block.kind == "table":
        return [" | ".join(row) for row in block.rows[:10]]
    return []


def _pptx_outline(title: str, blocks: list[ExportBlock]) -> list[tuple[str, list[str]]]:
    slides: list[tuple[str, list[str]]] = []
    current_title = title or "Orrery export"
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        lines = [line[:180] for line in current_lines if line.strip()]
        if lines or not slides:
            chunks = [lines[i:i + 8] for i in range(0, len(lines), 8)] or [[]]
            for index, chunk in enumerate(chunks):
                suffix = f" ({index + 1})" if index else ""
                slides.append(((current_title or title)[:80] + suffix, chunk))
        current_lines = []

    for block in blocks:
        if block.kind == "heading":
            if current_lines or slides:
                flush()
            current_title = _plain_spans(block.spans) or current_title
            continue
        current_lines.extend(_slide_text_from_block(block))
    if current_lines or not slides:
        flush()
    return slides[:60]


def _pptx_paragraph(line: str) -> str:
    bullet = line.startswith("- ")
    text = line[2:] if bullet else line
    ppr = '<a:pPr marL="342900" indent="-171450"><a:buChar char="•"/></a:pPr>' if bullet else ""
    return f'<a:p>{ppr}<a:r><a:rPr lang="en-US" sz="2000"/><a:t>{_xml(text)}</a:t></a:r></a:p>'


def _pptx_slide_xml(title: str, lines: list[str]) -> str:
    body = "".join(_pptx_paragraph(line) for line in lines) or '<a:p><a:r><a:rPr lang="en-US" sz="2000"/><a:t></a:t></a:r></a:p>'
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr>
      <p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>
      <p:sp>
        <p:nvSpPr><p:cNvPr id="2" name="Title"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr>
        <p:spPr><a:xfrm><a:off x="685800" y="457200"/><a:ext cx="10972800" cy="914400"/></a:xfrm></p:spPr>
        <p:txBody><a:bodyPr wrap="square"/><a:lstStyle/><a:p><a:r><a:rPr lang="en-US" sz="3400" b="1"/><a:t>{_xml(title)}</a:t></a:r></a:p></p:txBody>
      </p:sp>
      <p:sp>
        <p:nvSpPr><p:cNvPr id="3" name="Body"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr>
        <p:spPr><a:xfrm><a:off x="914400" y="1600200"/><a:ext cx="10363200" cy="4572000"/></a:xfrm></p:spPr>
        <p:txBody><a:bodyPr wrap="square"/><a:lstStyle/>{body}</p:txBody>
      </p:sp>
    </p:spTree>
  </p:cSld>
  <p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>
</p:sld>'''


def build_pptx(title: str, model: str, blocks: list[ExportBlock]) -> bytes:
    slides = _pptx_outline(title, blocks)
    now = datetime.datetime.now(datetime.timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
        slide_overrides = "".join(
            f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>'
            for i in range(1, len(slides) + 1)
        )
        zf.writestr("[Content_Types].xml", f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>
  <Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/>
  <Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/>
  <Override PartName="/ppt/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>
  {slide_overrides}
</Types>''')
        zf.writestr("_rels/.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>''')
        zf.writestr("docProps/core.xml", f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>{_xml(title)}</dc:title><dc:creator>Orrery</dc:creator><cp:lastModifiedBy>Orrery</cp:lastModifiedBy>
  <dcterms:created xsi:type="dcterms:W3CDTF">{now}</dcterms:created><dcterms:modified xsi:type="dcterms:W3CDTF">{now}</dcterms:modified>
</cp:coreProperties>''')
        zf.writestr("docProps/app.xml", f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>Orrery</Application><PresentationFormat>Wide</PresentationFormat><Slides>{len(slides)}</Slides>
</Properties>''')
        slide_ids = "".join(f'<p:sldId id="{255 + i}" r:id="rId{i + 1}"/>' for i in range(1, len(slides) + 1))
        zf.writestr("ppt/presentation.xml", f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:sldMasterIdLst><p:sldMasterId id="2147483648" r:id="rId1"/></p:sldMasterIdLst>
  <p:sldIdLst>{slide_ids}</p:sldIdLst>
  <p:sldSz cx="12192000" cy="6858000" type="wide"/><p:notesSz cx="6858000" cy="9144000"/>
</p:presentation>''')
        presentation_rels = ['<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="slideMasters/slideMaster1.xml"/>']
        for i in range(1, len(slides) + 1):
            presentation_rels.append(f'<Relationship Id="rId{i + 1}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{i}.xml"/>')
        zf.writestr("ppt/_rels/presentation.xml.rels", f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{''.join(presentation_rels)}</Relationships>''')
        zf.writestr("ppt/slideMasters/slideMaster1.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldMaster xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr></p:spTree></p:cSld><p:clrMap bg1="lt1" tx1="dk1" bg2="lt2" tx2="dk2" accent1="accent1" accent2="accent2" accent3="accent3" accent4="accent4" accent5="accent5" accent6="accent6" hlink="hlink" folHlink="folHlink"/><p:sldLayoutIdLst><p:sldLayoutId id="2147483649" r:id="rId1"/></p:sldLayoutIdLst><p:txStyles><p:titleStyle/><p:bodyStyle/><p:otherStyle/></p:txStyles></p:sldMaster>''')
        zf.writestr("ppt/slideMasters/_rels/slideMaster1.xml.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="../theme/theme1.xml"/></Relationships>''')
        zf.writestr("ppt/slideLayouts/slideLayout1.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldLayout xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" type="blank" preserve="1"><p:cSld name="Blank"><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr></p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr></p:sldLayout>''')
        zf.writestr("ppt/slideLayouts/_rels/slideLayout1.xml.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="../slideMasters/slideMaster1.xml"/></Relationships>''')
        zf.writestr("ppt/theme/theme1.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Orrery"><a:themeElements><a:clrScheme name="Orrery"><a:dk1><a:srgbClr val="0B1020"/></a:dk1><a:lt1><a:srgbClr val="FFFFFF"/></a:lt1><a:dk2><a:srgbClr val="172033"/></a:dk2><a:lt2><a:srgbClr val="E8ECF8"/></a:lt2><a:accent1><a:srgbClr val="F2B14E"/></a:accent1><a:accent2><a:srgbClr val="9DB9F0"/></a:accent2><a:accent3><a:srgbClr val="5BC489"/></a:accent3><a:accent4><a:srgbClr val="E36C6C"/></a:accent4><a:accent5><a:srgbClr val="8A94B8"/></a:accent5><a:accent6><a:srgbClr val="26314F"/></a:accent6><a:hlink><a:srgbClr val="3366CC"/></a:hlink><a:folHlink><a:srgbClr val="663399"/></a:folHlink></a:clrScheme><a:fontScheme name="Orrery"><a:majorFont><a:latin typeface="Arial"/></a:majorFont><a:minorFont><a:latin typeface="Arial"/></a:minorFont></a:fontScheme><a:fmtScheme name="Orrery"><a:fillStyleLst><a:solidFill><a:schemeClr val="lt1"/></a:solidFill></a:fillStyleLst><a:lnStyleLst><a:ln w="6350"><a:solidFill><a:schemeClr val="accent1"/></a:solidFill></a:ln></a:lnStyleLst><a:effectStyleLst><a:effectStyle><a:effectLst/></a:effectStyle></a:effectStyleLst><a:bgFillStyleLst><a:solidFill><a:schemeClr val="lt1"/></a:solidFill></a:bgFillStyleLst></a:fmtScheme></a:themeElements></a:theme>''')
        for i, (slide_title, slide_lines) in enumerate(slides, start=1):
            zf.writestr(f"ppt/slides/slide{i}.xml", _pptx_slide_xml(slide_title, slide_lines))
            zf.writestr(f"ppt/slides/_rels/slide{i}.xml.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/></Relationships>''')
    return output.getvalue()


def build_preview_html(title: str, model: str, content: str, export_format: str) -> bytes:
    if export_format == "html" and re.search(r"<!doctype html|<html[\s>]", content, re.I):
        return content.encode("utf-8")
    label = export_format.upper()
    body = _MARKDOWN.render(_clean_text(content)) if export_format != "txt" else f"<pre>{html.escape(_clean_text(content))}</pre>"
    document = f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>{html.escape(title)} preview</title><style>
body{{font-family:Arial,sans-serif;max-width:900px;margin:28px auto;padding:0 24px 48px;color:#172033;line-height:1.55}}
.preview-meta{{color:#69758a;font-size:12px;margin-bottom:20px}} table{{border-collapse:collapse;width:100%;margin:14px 0}}
th,td{{border:1px solid #c6cfdd;padding:7px 9px;text-align:left;vertical-align:top}} th{{background:#e8edf6}}
pre{{white-space:pre-wrap;background:#f3f5f9;border:1px solid #d7deea;border-radius:8px;padding:12px;overflow:auto}}
code{{font-family:Consolas,monospace}} blockquote{{border-left:4px solid #8fa5c8;margin-left:0;padding-left:14px;color:#4f5d73}}
</style></head><body><h1>{html.escape(title)}</h1><div class="preview-meta">{label} preview - Model: {html.escape(model or 'unknown')}</div>{body}</body></html>"""
    return document.encode("utf-8")

def render_export(title: str, model: str, content: str, export_format: str) -> ExportResult:
    if export_format not in SUPPORTED_FORMATS:
        raise ValueError("Unsupported export format")
    blocks = parse_markdown(content)
    slug = _slug(title)
    if export_format == "pdf":
        return ExportResult(build_pdf(title, model, blocks), "application/pdf", f"{slug}.pdf")
    if export_format == "docx":
        return ExportResult(
            build_docx(title, model, blocks),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{slug}.docx",
        )
    if export_format == "xlsx":
        return ExportResult(
            build_xlsx(title, model, blocks),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            f"{slug}.xlsx",
        )
    if export_format == "pptx":
        return ExportResult(
            build_pptx(title, model, blocks),
            "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            f"{slug}.pptx",
        )
    if export_format == "csv":
        return ExportResult(build_csv(blocks), "text/csv; charset=utf-8", f"{slug}.csv")
    if export_format == "md":
        return ExportResult(_plain_or_markdown_text(content, blocks, True), "text/markdown; charset=utf-8", f"{slug}.md")
    if export_format == "txt":
        return ExportResult(_plain_or_markdown_text(content, blocks, False), "text/plain; charset=utf-8", f"{slug}.txt")
    if export_format == "html":
        return ExportResult(build_html(title, model, content), "text/html; charset=utf-8", f"{slug}.html")
    if export_format == "json":
        return ExportResult((content.strip() + "\n").encode("utf-8"), "application/json; charset=utf-8", f"{slug}.json")
    raise ValueError("Unsupported export format")


def render_preview(title: str, model: str, content: str, export_format: str) -> ExportResult:
    if export_format == "pdf":
        return render_export(title, model, content, "pdf")
    return ExportResult(build_preview_html(title, model, content, export_format), "text/html; charset=utf-8", f"{_slug(title)}-preview.html")


async def _load_export_payload(conversation_id: str, message_id: str, export_format: str) -> tuple[str, str, str]:
    if export_format not in SUPPORTED_FORMATS:
        raise ValueError("Unsupported export format")
    try:
        conversation_uuid = uuid.UUID(conversation_id)
        message_uuid = uuid.UUID(message_id)
    except ValueError as exc:
        raise ExportNotFound("Reply not found") from exc

    async with get_sessionmaker()() as session:
        conversation = await session.get(Conversation, conversation_uuid)
        if conversation is None:
            raise ExportNotFound("Reply not found")
        messages = (
            await session.execute(
                select(Message).where(Message.conversation_id == conversation_uuid).order_by(Message.created_at)
            )
        ).scalars().all()

        target_index = -1
        for index, message in enumerate(messages):
            if message.id == message_uuid:
                target_index = index
                break
        if target_index < 0 or messages[target_index].role != "assistant":
            raise ExportNotFound("Reply not found")

        prompt = ""
        for previous in reversed(messages[:target_index]):
            if previous.role == "user":
                prompt = previous.context or previous.content
                break

        message = messages[target_index]
        title = conversation.title or "Orrery reply"
        model = message.model or conversation.model
        selected = select_export_content(prompt, message.content, export_format)
        return title, model, message.content, selected


async def export_message(conversation_id: str, message_id: str, export_format: str) -> ExportResult:
    from backend.features import docgen

    title, model, raw, selected = await _load_export_payload(conversation_id, message_id, export_format)
    spec = docgen.parse_doc_spec(raw)
    if spec is not None:  # model designed a structured file → build the real thing
        return await asyncio.to_thread(docgen.render_spec, title, model, spec, export_format)
    return await asyncio.to_thread(render_export, title, model, selected, export_format)


async def preview_message(conversation_id: str, message_id: str, export_format: str) -> ExportResult:
    from backend.features import docgen

    title, model, raw, selected = await _load_export_payload(conversation_id, message_id, export_format)
    spec = docgen.parse_doc_spec(raw)
    if spec is not None:
        return await asyncio.to_thread(docgen.render_spec_preview, title, model, spec, export_format)
    return await asyncio.to_thread(render_preview, title, model, selected, export_format)
